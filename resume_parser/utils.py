import io
import pdfplumber
from docx import Document

def extract_text_from_pdf(path_or_fileobj):
    text = []
    # pdfplumber supports file path or file-like object
    with pdfplumber.open(path_or_fileobj) as pdf:
        for page in pdf.pages:
            p = page.extract_text()
            if p:
                text.append(p)
    return "\n".join(text)

def extract_text_from_docx(path_or_fileobj):
    # python-docx expects a path or file-like (BytesIO) that it can read
    if hasattr(path_or_fileobj, 'read'):
        file_bytes = path_or_fileobj.read()
        doc = Document(io.BytesIO(file_bytes))
    else:
        doc = Document(path_or_fileobj)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)

def extract_text(file_field):
    filename = getattr(file_field, 'name', '')
    lower = filename.lower()
    # file_field may be an InMemoryUploadedFile; pass file_field.file or file_field
    f = file_field
    try:
        if lower.endswith('.pdf'):
            return extract_text_from_pdf(f)
        elif lower.endswith('.docx'):
            # ensure file pointer at start
            if hasattr(f, 'seek'):
                f.seek(0)
            return extract_text_from_docx(f)
        else:
            # fallback: try pdf then docx - or return empty
            if lower.endswith('.doc'):
                # doc (old format) needs special handling; recommend converting to docx
                return ""
            return ""
    except Exception as e:
        # log error in real app
        return ""
